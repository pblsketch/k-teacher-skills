#!/usr/bin/env python3
"""Production document-quality gate (open-safety + visual/usability).

This validator proves the individualized package's 12 artifacts are *actually
openable and usable* in a Korean classroom, not merely valid ZIP/XML:

  DOCX  -> reopened with python-docx (a production OOXML consumer): required
           parts (styles/settings/rels), >=1 section, A4 geometry, real tables
           and body paragraphs carrying the source content.
  HWPX  -> validated + reopened with python-hwpx (the official OWPML toolkit):
           standard required parts, package validation with zero ERROR issues,
           reopen ok, text/table extraction, residual fallback warnings captured
           and constrained to the known upstream manifest-optional-part set.
  HTML  -> print-ready: @page A4 + @media print rules, page containers, Korean
           font stack, audience-specific structure (teacher matrices / student
           answer space), self-contained (no external network dependency).

It intentionally FAILS against the pre-production minimal renderers (RED) and is
the GREEN gate for the production writers. Requires python-docx and python-hwpx
(see requirements-render.txt); import failure is a hard, honest error here.
"""
from __future__ import annotations

import io
import json
import logging
import re
import sys
import tempfile
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "tests"))

from renderers import render_package, extract_all, verify_parity  # noqa: E402
from validate_individualized_materials import (  # noqa: E402
    build_sample_ir,
    TEACHER_DOC,
    GROUP_DOCS,
)

# Production consumers. A missing dependency is a real failure of this gate.
import docx  # noqa: E402
from docx import Document  # noqa: E402
from docx.shared import Mm  # noqa: E402
import hwpx  # noqa: E402
from hwpx import HwpxDocument  # noqa: E402

# python-hwpx logs three fallbacks per reopen when the builder manifest omits the
# optional masterPage/history/version references. They are cosmetic (reopen ok,
# no ERROR issues) and originate entirely inside python-hwpx.
KNOWN_HWPX_FALLBACK_SUBSTRINGS = ("masterPage", "history", "version")


def assert_true(cond: bool, msg: str) -> None:
    if not cond:
        raise AssertionError(msg)


# --------------------------------------------------------------------------- #
# DOCX: real OOXML consumer reopen
# --------------------------------------------------------------------------- #

def check_docx(did: str, path: str) -> None:
    with zipfile.ZipFile(path) as z:
        names = set(z.namelist())
    for part in (
        "[Content_Types].xml",
        "word/document.xml",
        "word/styles.xml",
        "word/settings.xml",
        "word/_rels/document.xml.rels",
        "docProps/core.xml",
        "customXml/kteacher-backport-marker.json",
    ):
        assert_true(part in names, f"{did} docx: missing required OOXML part {part}")

    assert_true(Path(path).stat().st_size > 8000, f"{did} docx: production file too small ({Path(path).stat().st_size}B)")

    doc = Document(path)  # real python-docx reopen; raises if unreadable
    assert_true(len(doc.sections) >= 1, f"{did} docx: python-docx reopen has no section")
    sec = doc.sections[0]
    # A4 within 1mm tolerance.
    assert_true(abs(sec.page_width - Mm(210)) < Mm(1) and abs(sec.page_height - Mm(297)) < Mm(1),
                f"{did} docx: page geometry is not A4")
    body_texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert_true(len(body_texts) >= 4, f"{did} docx: too few body paragraphs ({len(body_texts)})")
    # A production plan/worksheet carries at least one real table (matrix / answer grid).
    assert_true(len(doc.tables) >= 1, f"{did} docx: no real table rendered")
    for t in doc.tables:
        assert_true(len(t.rows) >= 1 and len(t.columns) >= 1, f"{did} docx: degenerate table")
    # Named styles must exist (production styling, not a bare body).
    style_names = {s.name for s in doc.styles}
    assert_true("Title" in style_names or "Heading 1" in style_names, f"{did} docx: no heading/title styles")


# --------------------------------------------------------------------------- #
# HWPX: official OWPML toolkit validation + reopen
# --------------------------------------------------------------------------- #

def check_hwpx(did: str, path: str) -> list[str]:
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
    assert_true(names[0] == "mimetype" and z_read_mimetype(path) == b"application/hwp+zip",
                f"{did} hwpx: OCF mimetype must be first and stored")
    nameset = set(names)
    for part in (
        "version.xml",
        "Contents/header.xml",
        "Contents/section0.xml",
        "settings.xml",
        "Preview/PrvText.txt",
        "META-INF/container.xml",
        "META-INF/manifest.xml",
        "META-INF/container.rdf",
        "Contents/content.hpf",
        "META-INF/kteacher-backport-marker.json",
    ):
        assert_true(part in nameset, f"{did} hwpx: missing standard part {part}")

    report = hwpx.validate_package(path)
    errors = [i for i in report.issues if i.level == "error"]
    assert_true(not errors, f"{did} hwpx: package validation ERROR issues: {[i.message for i in errors]}")

    # Reopen with python-hwpx, capturing residual library warnings.
    logbuf = io.StringIO()
    handler = logging.StreamHandler(logbuf)
    hwlog = logging.getLogger("hwpx")
    prev_level = hwlog.level
    hwlog.addHandler(handler)
    hwlog.setLevel(logging.WARNING)
    try:
        doc = HwpxDocument.open(path)
        text = doc.export_text()
    finally:
        hwlog.removeHandler(handler)
        hwlog.setLevel(prev_level)

    assert_true(text.strip(), f"{did} hwpx: reopened document extracted no text")
    warnings = [ln for ln in logbuf.getvalue().splitlines() if ln.strip()]
    for w in warnings:
        assert_true(any(sub in w for sub in KNOWN_HWPX_FALLBACK_SUBSTRINGS),
                    f"{did} hwpx: unexpected reopen warning: {w}")
    return warnings


def z_read_mimetype(path: str) -> bytes:
    with zipfile.ZipFile(path) as z:
        return z.read("mimetype")


# --------------------------------------------------------------------------- #
# HTML: print-ready + audience-specific + self-contained
# --------------------------------------------------------------------------- #

def check_html(did: str, path: str, *, is_teacher: bool) -> None:
    text = Path(path).read_text(encoding="utf-8")
    assert_true("@page" in text and "A4" in text, f"{did} html: no @page A4 rule")
    assert_true("@media print" in text, f"{did} html: no print stylesheet")
    assert_true(re.search(r"Pretendard|Apple SD Gothic|Malgun|Noto Sans KR|system-ui", text),
                f"{did} html: no Korean/system font stack")
    # Multi-page A4 composition: teacher plan is 3 physical pages, each worksheet 2.
    pages = re.findall(r'<section class="page"', text)
    expected_pages = 3 if is_teacher else 2
    assert_true(len(pages) == expected_pages,
                f"{did} html: expected {expected_pages} A4 <section class=\"page\">, got {len(pages)}")
    assert_true("break-after:page" in text or "break-before:page" in text,
                f"{did} html: no explicit print page break between pages")
    # Self-contained: no external network dependency.
    assert_true(not re.search(r'(src|href)\s*=\s*"https?://', text), f"{did} html: external network dependency present")
    assert_true("<h2" in text, f"{did} html: no section hierarchy (h2)")
    # Every doc carries at least one real table (matrix / grid / answer space).
    assert_true(text.count("<table") >= 2, f"{did} html: too few structured tables ({text.count('<table')})")
    if is_teacher:
        assert_true("45분 수업 흐름" in text and "모둠별 배치 비교" in text,
                    f"{did} html: teacher plan missing lesson-flow / A-B-C deployment matrix")
    else:
        assert_true(re.search(r"answer|답|서술|칸", text), f"{did} html: student worksheet has no answer space")
        for ph in ("학급 제공 자료", "그림과 표로 정보를 함께 제시한다"):
            assert_true(ph not in text, f"{did} html: student worksheet contains generic placeholder {ph!r}")


# --------------------------------------------------------------------------- #

def main() -> None:
    ir = build_sample_ir()
    all_warnings: dict[str, list[str]] = {}
    with tempfile.TemporaryDirectory() as td:
        rendered = render_package(ir, td)
        assert_true(len(rendered) == 4, f"expected 4 documents, got {len(rendered)}")
        files = [p for doc in rendered.values() for p in doc.values()]
        assert_true(len(files) == 12, f"expected 12 files, got {len(files)}")

        for did, paths in rendered.items():
            is_teacher = did == TEACHER_DOC
            check_docx(did, paths["docx"])
            all_warnings[did] = check_hwpx(did, paths["hwpx"])
            check_html(did, paths["html"], is_teacher=is_teacher)

            # Marker + semantic round-trip/parity must survive the production writers.
            ex = extract_all(paths)
            ok, reasons = verify_parity(ex)
            assert_true(ok, f"{did}: 3-format parity failed after production render: {reasons}")

    print("PASS validate_production_document_quality")
    print("- DOCX reopened by python-docx: styles/settings/rels, A4, sections>=1, real tables + paragraphs")
    print("- HWPX validated + reopened by python-hwpx: standard parts, 0 ERROR issues, text extracted")
    print("- HWPX residual warnings constrained to known upstream manifest fallbacks:")
    for did, ws in all_warnings.items():
        print(f"    {did}: {len(ws)} warning line(s)")
    print("- HTML print-ready: @page A4, @media print, Korean font stack, self-contained, audience-specific")
    print("- marker + 3-format semantic parity preserved through production writers")


if __name__ == "__main__":
    try:
        main()
    except AssertionError as error:
        print(f"FAIL validate_production_document_quality: {error}")
        raise SystemExit(1) from error
