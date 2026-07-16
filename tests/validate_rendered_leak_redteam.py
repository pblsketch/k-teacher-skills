#!/usr/bin/env python3
"""Red-team the *rendered* student artifacts (not just the IR): no teacher-only
diagnostic term or pathway-profile label may surface in any student DOCX/HWPX/HTML
or its content sidecar — including nested block keys/values and separator evasion.
Includes a positive control (a planted term is caught) and a false-positive control
(legitimate subject vocabulary is not flagged).
"""
from __future__ import annotations

import copy
import re
import sys
import tempfile
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "tests"))

from renderers import render_package  # noqa: E402
from providers.materials import worksheet as ws  # noqa: E402
from providers.materials.individualized import STUDENT_DIAGNOSTIC_TERMS  # noqa: E402
from docx import Document  # noqa: E402
from hwpx import HwpxDocument  # noqa: E402
from validate_individualized_materials import (  # noqa: E402
    build_sample_ir, sample_pathways, GROUP_DOCS,
)

_TAG = re.compile(r"<[^>]+>")


def assert_true(cond: bool, msg: str) -> None:
    if not cond:
        raise AssertionError(msg)


def student_text(paths: dict) -> str:
    """All human-visible + sidecar text recoverable from a rendered student doc."""
    chunks: list[str] = []
    d = Document(paths["docx"])
    chunks += [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    chunks.append(HwpxDocument.open(paths["hwpx"]).export_text())
    html = Path(paths["html"]).read_text(encoding="utf-8")
    chunks.append(_TAG.sub(" ", html))
    # content sidecars (round-trip encoding, incl. data-block-json)
    with zipfile.ZipFile(paths["docx"]) as z:
        chunks.append(z.read("customXml/kteacher-content.xml").decode("utf-8"))
    with zipfile.ZipFile(paths["hwpx"]) as z:
        chunks.append(z.read("Contents/kteacher-content.xml").decode("utf-8"))
    return "\n".join(chunks)


def main() -> None:
    ir = build_sample_ir()
    profile_labels = [p.teacher_profile_label for p in sample_pathways()]

    with tempfile.TemporaryDirectory() as td:
        rendered = render_package(ir, td)
        for gid in GROUP_DOCS.values():
            text = student_text(rendered[gid])
            for term in STUDENT_DIAGNOSTIC_TERMS:
                assert_true(not ws.contains_forbidden_term(text, term),
                            f"{gid}: diagnostic term '{term}' leaked into rendered student artifact")
            for label in profile_labels:
                assert_true(label not in text,
                            f"{gid}: teacher profile label leaked into rendered student artifact")

        # Positive control: a planted diagnostic term must be detectable in the render.
        planted = copy.deepcopy(ir)
        sdoc = next(d for d in planted["lesson_package"]["documents"] if d["document_id"] == GROUP_DOCS["Group A"])
        sdoc["content"]["blocks"].append(ws.student_note("leak-probe", text="이 학생은 기초그룹 학생입니다."))
        pr = render_package(planted, Path(td) / "planted")
        assert_true(ws.contains_forbidden_term(student_text(pr[GROUP_DOCS["Group A"]]), "기초그룹"),
                    "positive control failed: planted diagnostic term was not detectable in render")

        # False-positive control: legitimate subject vocabulary must NOT trip the scan.
        for phrase in ("표준편차를 구하시오.", "기초 대사량을 계산한다.", "발표 준비를 한다."):
            assert_true(all(not ws.contains_forbidden_term(phrase, term) for term in STUDENT_DIAGNOSTIC_TERMS),
                        f"false-positive control failed: legitimate phrase flagged: {phrase}")

    print("PASS validate_rendered_leak_redteam")
    print("- no diagnostic term / pathway label surfaces in any rendered student DOCX/HWPX/HTML or sidecar")
    print("- positive control: a planted diagnostic term is caught in the actual render")
    print("- false-positive control: 표준편차 / 기초 대사량 / 발표 준비 are not flagged")


if __name__ == "__main__":
    try:
        main()
    except AssertionError as error:
        print(f"FAIL validate_rendered_leak_redteam: {error}")
        raise SystemExit(1) from error
