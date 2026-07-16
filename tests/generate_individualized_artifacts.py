#!/usr/bin/env python3
"""Generate the 12 individualized-package artifacts + a delivery ZIP from the
committed golden fixture, then capture real production-consumer reopen evidence.

Outputs (all under the ignored artifacts/ tree; nothing is committed):
  artifacts/individualized/<doc>.{hwpx,docx,html}   -- 4 docs x 3 formats = 12
  artifacts/individualized/delivery.zip             -- one delivery bundle
  .gjc/qa/evidence/GREEN-reopen-evidence.txt        -- reopen receipts
"""
from __future__ import annotations

import io
import json
import logging
import sys
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "tests"))

from renderers import render_package, extract_all, verify_parity  # noqa: E402
from providers.materials.individualized import validate_individualized_package  # noqa: E402
from validate_individualized_materials import build_sample_ir, TEACHER_DOC  # noqa: E402

import docx  # noqa: E402
from docx import Document  # noqa: E402
from docx.shared import Mm  # noqa: E402
import hwpx  # noqa: E402
from hwpx import HwpxDocument  # noqa: E402

OUT = ROOT / "artifacts" / "individualized"
EVID = ROOT / ".gjc" / "qa" / "evidence"


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    EVID.mkdir(parents=True, exist_ok=True)

    ir = build_sample_ir()
    golden = json.loads((ROOT / "tests" / "golden" / "individualized-materials" / "9과17-01.package.json").read_text(encoding="utf-8"))
    assert ir == golden, "sample IR drifted from committed golden"
    assert validate_individualized_package(ir) == [], "golden package must be contract-clean"

    rendered = render_package(ir, OUT)
    lines: list[str] = ["# Production reopen / render evidence", ""]

    files: list[Path] = []
    for did, paths in rendered.items():
        lines.append(f"## {did}")
        # DOCX: python-docx reopen
        d = Document(paths["docx"])
        sec = d.sections[0]
        a4 = abs(sec.page_width - Mm(210)) < Mm(1) and abs(sec.page_height - Mm(297)) < Mm(1)
        with zipfile.ZipFile(paths["docx"]) as z:
            dparts = len(z.namelist())
        lines.append(
            f"- DOCX  {Path(paths['docx']).stat().st_size:>6}B parts={dparts} "
            f"reopen(python-docx) sections={len(d.sections)} A4={a4} "
            f"paras={sum(1 for p in d.paragraphs if p.text.strip())} tables={len(d.tables)}"
        )
        # HWPX: validate + reopen (capture library warnings)
        rep = hwpx.validate_package(paths["hwpx"])
        errs = [i for i in rep.issues if i.level == "error"]
        warns = [i for i in rep.issues if i.level == "warning"]
        logbuf = io.StringIO()
        handler = logging.StreamHandler(logbuf)
        hwlog = logging.getLogger("hwpx")
        hwlog.addHandler(handler)
        hwlog.setLevel(logging.WARNING)
        try:
            hd = HwpxDocument.open(paths["hwpx"])
            text = hd.export_text()
        finally:
            hwlog.removeHandler(handler)
        reopen_warn_lines = [ln for ln in logbuf.getvalue().splitlines() if ln.strip()]
        with zipfile.ZipFile(paths["hwpx"]) as z:
            hparts = len(z.namelist())
        lines.append(
            f"- HWPX  {Path(paths['hwpx']).stat().st_size:>6}B parts={hparts} "
            f"validate(errors={len(errs)}, warnings={len(warns)}) reopen(python-hwpx)=ok "
            f"chars={len(text.strip())} reopen_warn_lines={len(reopen_warn_lines)}"
        )
        for w in warns:
            lines.append(f"    validate-warning: {w.message}")
        for w in reopen_warn_lines:
            lines.append(f"    reopen-warning: {w.strip()}")
        # HTML
        htext = Path(paths["html"]).read_text(encoding="utf-8")
        lines.append(
            f"- HTML  {Path(paths['html']).stat().st_size:>6}B @page-A4={'@page' in htext and 'A4' in htext} "
            f"print-css={'@media print' in htext} self-contained={'http://' not in htext and 'https://' not in htext}"
        )
        # parity
        ok, reasons = verify_parity(extract_all(paths))
        lines.append(f"- PARITY 3-format={'ok' if ok else reasons}")
        lines.append("")
        files.extend(Path(p) for p in paths.values())

    assert len(files) == 12, f"expected 12 files, got {len(files)}"

    delivery = OUT / "delivery.zip"
    with zipfile.ZipFile(delivery, "w", zipfile.ZIP_DEFLATED) as z:
        for f in files:
            z.write(f, arcname=f.name)
    lines.append(f"Delivery bundle: {delivery.relative_to(ROOT)} ({delivery.stat().st_size}B, {len(files)} files)")

    report = "\n".join(lines) + "\n"
    (EVID / "GREEN-reopen-evidence.txt").write_text(report, encoding="utf-8")
    print(report)
    print(f"12 artifacts + delivery.zip under {OUT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
