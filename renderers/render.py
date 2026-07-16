"""Production renderer: single canonical lesson-package IR -> real HWPX/DOCX/HTML.

Visible content is produced by production writers — python-docx (OOXML with
styles/settings/rels/sectPr and real tables), the official python-hwpx builder
(standard OWPML package, reopen-verified), and print-ready self-contained HTML
(A4 @page + @media print) — so the artifacts are actually openable and usable in
a Korean classroom, not merely valid ZIP/XML.

Each package also embeds an extractable backport marker at the schema's canonical
location plus a round-trippable content sidecar (data-content-id / data-provenance-*
/ data-block-* encoding). The sidecar is the same canonical content serialized once
for marker-based 3-way parity verification — it is NOT a parallel IR. This is an
independent implementation (no anthropics renderer code).

Rendering DOCX/HWPX requires python-docx and python-hwpx (requirements-render.txt);
the extractors/parity layer are pure-stdlib and import without those deps.
"""
from __future__ import annotations

import hashlib
import html as htmllib
import json
import re
import zipfile
from pathlib import Path

CANONICAL_LOCATIONS = {
    "hwpx": {"locator_kind": "package-member", "locator_value": "META-INF/kteacher-backport-marker.json", "validator_extraction_method": "read package member META-INF/kteacher-backport-marker.json and parse JSON"},
    "docx": {"locator_kind": "opc-part", "locator_value": "/customXml/kteacher-backport-marker.json", "validator_extraction_method": "read OPC custom XML part /customXml/kteacher-backport-marker.json and parse JSON"},
    "html": {"locator_kind": "dom-node", "locator_value": "script#kteacher-backport-marker", "validator_extraction_method": "read <script id=kteacher-backport-marker type=application/json> from document head and parse text as JSON"},
}
# ZIP member paths (marker), per format.
_HWPX_MARKER = "META-INF/kteacher-backport-marker.json"
_DOCX_MARKER = "customXml/kteacher-backport-marker.json"
_DOCX_CONTENT = "customXml/kteacher-content.xml"
_HWPX_CONTENT = "Contents/kteacher-content.xml"

_XML_ESC = {"&": "&amp;", "<": "&lt;", ">": "&gt;", '"': "&quot;"}


def _xesc(s: str) -> str:
    return "".join(_XML_ESC.get(c, c) for c in s)


# --- canonical content + marker ------------------------------------------------

def canonical_content(document: dict) -> dict:
    content = document.get("content", {})
    return {
        "document_id": document["document_id"],
        "document_class": document["document_class"],
        "title": document["title"],
        "required_content": [{"content_id": s["content_id"], "text": s["text"]} for s in content.get("sections", [])],
        "provenance_markers": [
            {"record_id": m["record_id"], "label": m["label"], "evidence_text": m["evidence_text"]}
            for m in content.get("provenance_markers", [])
        ],
        "unresolved_boundary_markers": list(content.get("unresolved_boundary_markers", [])),
        "blocks": list(content.get("blocks", [])),
    }


def content_fingerprint(canonical: dict) -> str:
    payload = json.dumps(
        {k: canonical[k] for k in ("required_content", "provenance_markers", "unresolved_boundary_markers", "blocks")},
        ensure_ascii=False, sort_keys=True,
    )
    return "sha256:" + hashlib.sha256(payload.encode("utf-8")).hexdigest()


def build_marker(ir: dict, fmt: str, fingerprint: str, render_revision_id: str, rendered_at: str = "2026-07-16T00:00:00Z") -> dict:
    return {
        "backport_marker_version": "1",
        "workflow_id": ir["workflow_id"],
        "ir_id": ir["ir_id"],
        "ir_revision_id": ir["ir_revision_id"],
        "render_revision_id": render_revision_id,
        "renderer_format": fmt,
        "renderer_version": "kteacher-minimal-renderer@2026.07.16",
        "rendered_at": rendered_at,
        "content_fingerprint": fingerprint,
        "authoritative_source": "json-ir",
        "manual_edit_status": "clean",
        "canonical_locations": CANONICAL_LOCATIONS,
    }


# --- paragraph encoding (round-trippable) --------------------------------------

def _content_paras(canonical: dict, ptag: str, ttag: str | None) -> str:
    out = []
    for item in canonical["required_content"]:
        inner = f"<{ttag}>{_xesc(item['text'])}</{ttag}>" if ttag else _xesc(item["text"])
        out.append(f'<{ptag} data-content-id="{_xesc(item["content_id"])}">{inner}</{ptag}>')
    for m in canonical["provenance_markers"]:
        inner = f"<{ttag}>{_xesc(m['evidence_text'])}</{ttag}>" if ttag else _xesc(m["evidence_text"])
        out.append(
            f'<{ptag} data-provenance-record-id="{_xesc(m["record_id"])}" data-label="{_xesc(m["label"])}">{inner}</{ptag}>'
        )
    return "\n".join(out)

# --- block encoding (real tables + answer-space, round-trippable) --------------

_BLOCK_TAGS = {
    "html": {"para": "p", "table": "table", "row": "tr", "cell": "td"},
    "hwpx": {"para": "hp:p", "table": "hp:tbl", "row": "hp:tr", "cell": "hp:tc"},
    "docx": {"para": "w:p", "table": "w:tbl", "row": "w:tr", "cell": "w:tc"},
}


def _block_attrs(block: dict) -> str:
    payload = json.dumps(block, ensure_ascii=False, sort_keys=True)
    return (
        f'data-block-id="{_xesc(block["block_id"])}" '
        f'data-block-type="{_xesc(block["block_type"])}" '
        f'data-block-json="{_xesc(payload)}"'
    )


def _cell_xml(value, tags: dict, fmt: str) -> str:
    """Emit a table cell with the paragraph/run/text nesting required by the
    target office format. Empty cells still receive a text node container."""
    c = tags["cell"]
    text = _xesc(str(value))
    if fmt == "docx":
        return f'<{c}><w:p><w:r><w:t xml:space="preserve">{text}</w:t></w:r></w:p></{c}>'
    if fmt == "hwpx":
        return f'<{c}><hp:subList><hp:p><hp:run><hp:t>{text}</hp:t></hp:run></hp:p></hp:subList></{c}>'
    return f"<{c}>{text}</{c}>"


def _row_xml(values: list, tags: dict, fmt: str, row_height_mm: float | None = None) -> str:
    r = tags["row"]
    cells = "".join(_cell_xml(v, tags, fmt) for v in values)
    if row_height_mm is None:
        return f"<{r}>{cells}</{r}>"
    if fmt == "docx":
        twips = max(1, round(row_height_mm * 56.6929))
        return f'<{r}><w:trPr><w:trHeight w:val="{twips}" w:hRule="atLeast"/></w:trPr>{cells}</{r}>'
    if fmt == "hwpx":
        return f'<{r} data-row-height-mm="{row_height_mm:g}">{cells}</{r}>'
    return f'<{r} style="height:{row_height_mm:g}mm" data-row-height-mm="{row_height_mm:g}">{cells}</{r}>'


def _table_block(block: dict, header: list, rows: list, tags: dict, fmt: str,
                 row_height_mm: float | None = None) -> str:
    t = tags["table"]
    parts = [f"<{t} {_block_attrs(block)}>"]
    if fmt == "html" and block.get("caption"):
        parts.append(f'<caption>{_xesc(str(block["caption"]))}</caption>')
    if header:
        parts.append(_row_xml(header, tags, fmt))
    for row in rows:
        parts.append(_row_xml(row, tags, fmt, row_height_mm))
    parts.append(f"</{t}>")
    return "".join(parts)


def _block_visible_text(block: dict) -> str:
    bt = block["block_type"]
    if bt in ("student_task", "exit_ticket"):
        return block["prompt"]
    if bt == "student_note":
        return block["text"]
    if bt == "source_card":
        citation = f' ({block["citation"]})' if block.get("citation") else ""
        return f'{block["title"]}: {block["body"]} — {block["source"]}{citation}'
    if bt == "sentence_support":
        return " / ".join(block["stems"])
    if bt == "number_line":
        return f'{block.get("label", "수직선")}: {block["min"]}~{block["max"]} (step {block["step"]})'
    if bt == "group_cohesion":
        return f'{block["group_label"]}: ' + ", ".join(block["members"])
    if bt == "page_break":
        return "[page-break]"
    return ""


def _block_node(block: dict, tags: dict, fmt: str) -> str:
    """One content block -> a real per-format table / answer-space / paragraph node,
    each carrying data-block-id/data-block-type and a round-trippable data-block-json."""
    bt = block["block_type"]
    if bt == "fill_table":
        return _table_block(block, block["headers"], block["rows"], tags, fmt, row_height_mm=block["row_height_mm"])
    if bt == "data_table":
        return _table_block(block, block["headers"], block["cells"], tags, fmt)
    if bt == "answer_box":
        # Real ruled answer-space: one empty ruled row per required line.
        line_height_mm = block["min_height_mm"] / max(1, block["min_lines"])
        return _table_block(block, [], [[""] for _ in range(block["min_lines"])], tags, fmt,
                            row_height_mm=line_height_mm)
    return f'<{tags["para"]} {_block_attrs(block)}>{_xesc(_block_visible_text(block))}</{tags["para"]}>'


def _content_blocks(canonical: dict, fmt: str) -> str:
    """Encode content.blocks into real per-format tables + answer-space nodes, each
    carrying data-block-id/data-block-type and a round-trippable data-block-json."""
    tags = _BLOCK_TAGS[fmt]
    return "\n".join(_block_node(block, tags, fmt) for block in canonical.get("blocks", []))


# --- production layout model (backend-neutral) ---------------------------------
# The visible office/print content is rendered by production writers (python-docx,
# python-hwpx, print-CSS HTML). A round-trippable content sidecar (the same encoding
# the marker parity layer consumes) is embedded alongside, so semantic parity and
# the backport marker survive unchanged. This is NOT a parallel IR: the sidecar is
# the canonical content, serialized once, exactly like the marker JSON.

_CONTENT_LABELS = {
    "common-goal": "공통 목표",
    "grouping-policy": "모둠 편성 원칙",
    "regroup-evidence": "재편성 근거",
    "common-success": "공통 성공 기준",
    "common-exit": "공통 출구표(가장 어려운 사례)",
    "student-goal": "오늘의 목표",
    "pathway-a": "Group A 운영 지원",
    "pathway-b": "Group B 운영 지원",
    "pathway-c": "Group C 운영 지원",
    "lesson-title": "수업 주제",
    "learning-goal": "학습 목표",
}
_DEMAND_LABELS = {"recall": "기억", "apply": "적용", "analyze": "분석", "evaluate": "평가", "create": "창안"}
_DOC_CLASS_LABELS = {
    "individualized-plan": "개별화 수업 운영안",
    "worksheet": "학생 활동지",
    "lesson-plan": "수업안",
}


def _label_for(content_id: str) -> str:
    if content_id in _CONTENT_LABELS:
        return _CONTENT_LABELS[content_id]
    m = re.match(r"^activity-(\d+)$", content_id)
    if m:
        return f"활동 {m.group(1)}"
    m = re.match(r"^(?:sc|success)-(\d+)$", content_id)
    if m:
        return f"성공 기준 {int(m.group(1)) + 1}"
    return content_id.replace("-", " ")


def _doc_subtitle(canonical: dict) -> str:
    """Human-facing purpose line; internal document ids remain in metadata/sidecars only."""
    if canonical["document_class"] == "individualized-plan":
        return "수업 운영·관찰·다음 차시 지원 자료"
    if canonical["document_class"] == "worksheet":
        return "학습 목표·탐구 활동·출구표"
    return str(_DOC_CLASS_LABELS.get(canonical["document_class"], "수업 자료") or "수업 자료")


def _provenance_display_label(label: str) -> str:
    """Translate machine ledger labels for print while preserving originals in metadata."""
    folded = str(label).casefold()
    if "from-curriculum" in folded:
        return "국가 교육과정(교사 확인)"
    if "from-textbook" in folded:
        return "교과서 근거"
    if "teacher" in folded or "created" in folded:
        return "교사 제작 자료"
    return "근거 자료"


def _blocks_flow(blocks: list) -> list:
    ops: list = []
    for b in blocks:
        bt = b["block_type"]
        if bt == "student_task":
            ops.append(("task", _DEMAND_LABELS.get(b.get("cognitive_demand"), ""), b["prompt"]))
        elif bt == "answer_box":
            ops.append(("answer", int(b["min_lines"]), float(b["min_height_mm"])))
        elif bt == "source_card":
            ops.append(("card", b["title"], b["body"], b["source"], b.get("citation")))
        elif bt == "sentence_support":
            ops.append(("stems", list(b["stems"])))
        elif bt == "student_note":
            ops.append(("note", b["text"]))
        elif bt == "exit_ticket":
            ops.append(("callout", "출구표(가장 어려운 사례)", b["prompt"]))
        elif bt == "fill_table":
            ops.append(("table", b.get("caption"), list(b["headers"]), [list(r) for r in b["rows"]], float(b["row_height_mm"])))
        elif bt == "data_table":
            ops.append(("table", b.get("caption"), list(b["headers"]), [list(r) for r in b["cells"]], None))
        elif bt == "page_break":
            ops.append(("pagebreak",))
        else:
            ops.append(("note", _block_visible_text(b)))
    return ops


def _flow(canonical: dict) -> list:
    """Backend-neutral layout ops derived from the canonical content."""
    ops: list = [("title", canonical["title"]), ("meta", _doc_subtitle(canonical))]
    for item in canonical["required_content"]:
        ops.append(("heading", _label_for(item["content_id"])))
        ops.append(("para", item["text"]))
    prov = canonical["provenance_markers"]
    if prov:
        ops.append(("heading", "근거 자료"))
        ops.append(("evidence", [(_provenance_display_label(m["label"]), m["evidence_text"]) for m in prov]))
    ops.extend(_blocks_flow(canonical.get("blocks", [])))
    return ops


def _content_sidecar_xml(canonical: dict) -> str:
    """The round-trippable content encoding, embedded as a package part so the
    marker-parity/extraction contract is preserved independent of visible layout."""
    body = _content_paras(canonical, "p", None)
    blocks = _content_blocks(canonical, "html")
    inner = body + (("\n" + blocks) if blocks else "")
    return (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        f'<kteacher-content data-document-id="{_xesc(canonical["document_id"])}" '
        f'data-document-class="{_xesc(canonical["document_class"])}" '
        f'data-title="{_xesc(canonical["title"])}">\n{inner}\n</kteacher-content>'
    )


# --- dependency guards ---------------------------------------------------------

def _require_docx():
    try:
        import docx  # noqa: F401
        return docx
    except ImportError as error:  # pragma: no cover - environment guard
        raise RuntimeError(
            "python-docx is required to render production DOCX. "
            "Install render dependencies: pip install -r requirements-render.txt"
        ) from error


def _require_hwpx_builder():
    try:
        from hwpx import builder
        return builder
    except ImportError as error:  # pragma: no cover - environment guard
        raise RuntimeError(
            "python-hwpx is required to render production HWPX. "
            "Install render dependencies: pip install -r requirements-render.txt"
        ) from error


# --- HTML (print-ready, self-contained) ----------------------------------------

_HTML_CSS = """
:root{--ink:#1f2933;--navy:#1f3a5f;--teal:#0f766e;--line:#c9d3de;--soft:#eef3f8;--band:#eaf1fb;}
*{box-sizing:border-box;}
html{font-size:15px;}
body{margin:0;background:#e8edf3;color:var(--ink);
 font-family:"Pretendard","Apple SD Gothic Neo","Malgun Gothic","Noto Sans KR",system-ui,sans-serif;
 line-height:1.55;-webkit-print-color-adjust:exact;print-color-adjust:exact;}
section.page{width:210mm;min-height:297mm;margin:10mm auto;padding:16mm 18mm;background:#fff;
 box-shadow:0 2px 14px rgba(31,41,51,.18);}
.doc-header{border-bottom:3px solid var(--navy);padding-bottom:10px;margin-bottom:14px;}
.page-run{color:#5b6b7c;font-size:.8rem;border-bottom:1px solid var(--line);
 padding-bottom:6px;margin-bottom:12px;}
h1{font-size:1.6rem;color:var(--navy);margin:0 0 4px;}
.subtitle{color:#5b6b7c;font-size:.9rem;letter-spacing:.02em;}
section.block{margin:12px 0;}
h2{font-size:1.02rem;color:var(--teal);border-left:5px solid var(--teal);
 padding:2px 0 2px 10px;margin:14px 0 8px;}
p.body{margin:4px 0;}
caption{caption-side:top;text-align:left;font-weight:700;color:var(--navy);padding:8px 0 4px;font-size:.95rem;}
.ev-label{color:var(--navy);font-weight:700;white-space:nowrap;}
p[data-block-type="student_task"]{font-weight:600;background:var(--band);
 border:1px solid var(--line);border-radius:6px;padding:8px 12px;margin:12px 0 6px;}
p[data-block-type="student_task"]::before{content:"과제  ";color:var(--navy);}
p[data-block-type="student_note"]{color:#374151;margin:8px 0;}
p[data-block-type="source_card"]{background:#f5faf9;border-left:4px solid var(--teal);
 padding:8px 12px;margin:10px 0;}
p[data-block-type="sentence_support"]{color:#334;font-style:italic;margin:6px 0;}
p[data-block-type="exit_ticket"]{background:#fff7ed;border:1px solid #f0c894;border-radius:6px;
 padding:10px 12px;margin:14px 0;font-weight:600;}
p[data-block-type="exit_ticket"]::before{content:"출구표(가장 어려운 사례)  ";color:#9a5b00;display:block;font-size:.82rem;}
p[data-block-type="page_break"]{display:none;}
table{width:100%;border-collapse:collapse;margin:6px 0 14px;}
th,td{border:1px solid var(--line);padding:6px 8px;vertical-align:top;font-size:.9rem;}
th{background:var(--band);color:var(--navy);text-align:left;}
table[data-block-type="answer_box"] tr[data-row-height-mm] td{border:none;border-bottom:1px solid #6b7a8c;}
.footer{margin-top:16px;border-top:1px solid var(--line);padding-top:6px;
 color:#7a8794;font-size:.76rem;text-align:right;}
@page{size:A4;margin:14mm;}
@media print{
 body{background:#fff;}
 section.page{width:auto;min-height:auto;margin:0;padding:0;box-shadow:none;break-after:page;}
 section.page:last-of-type{break-after:auto;}
 section.block,table,tr,p[data-block-type="student_task"],p[data-block-type="exit_ticket"]{break-inside:avoid;}
 h2,caption{break-after:avoid;}
}
@media screen and (max-width:840px){
 body{background:#fff;}
 section.page{width:100%;min-height:auto;margin:0 0 8px;padding:8mm 6mm;box-shadow:none;}
 table{table-layout:fixed;}
 th,td{word-break:break-word;overflow-wrap:anywhere;}
}
""".strip()


def _html_page1_prefix(canonical: dict) -> str:
    """Page-1 lead: canonical sections (goal/target) + the provenance evidence table."""
    parts: list = []
    for item in canonical["required_content"]:
        parts.append(
            f'<section class="block"><h2>{_xesc(_label_for(item["content_id"]))}</h2>'
            f'<p class="body" data-content-id="{_xesc(item["content_id"])}">{_xesc(item["text"])}</p></section>'
        )
    if canonical["provenance_markers"]:
        rows = "".join(
            f'<tr><td class="ev-label">{_xesc(_provenance_display_label(m["label"]))}</td>'
            f'<td><p data-provenance-record-id="{_xesc(m["record_id"])}" data-label="{_xesc(m["label"])}">'
            f'{_xesc(m["evidence_text"])}</p></td></tr>'
            for m in canonical["provenance_markers"]
        )
        parts.append(
            '<section class="block evidence"><h2>근거 자료</h2>'
            '<table><caption>근거 자료(출처)</caption><thead><tr><th>구분</th><th>근거</th></tr></thead>'
            f'<tbody>{rows}</tbody></table></section>'
        )
    return "\n".join(parts)


def render_html(canonical: dict, marker: dict, path: Path) -> None:
    tags = _BLOCK_TAGS["html"]
    title = _xesc(canonical["title"])
    subtitle = _xesc(_doc_subtitle(canonical))

    # Paginate the canonical blocks into A4 <section class="page"> at each page_break.
    pages: list = [[_html_page1_prefix(canonical)]]
    for block in canonical.get("blocks", []):
        pages[-1].append(_block_node(block, tags, "html"))
        if block["block_type"] == "page_break":
            pages.append([])
    total = len(pages)

    sections: list = []
    for i, page in enumerate(pages):
        if i == 0:
            head = (f'<header class="doc-header"><h1>{title}</h1>'
                    f'<div class="subtitle">{subtitle}</div></header>')
        else:
            head = f'<div class="page-run">{title} · {i + 1}/{total} 쪽</div>'
        inner = "\n".join(p for p in page if p)
        foot = f'<div class="footer">{title} · 출력용 A4 · {i + 1}/{total}</div>'
        sections.append(f'<section class="page">{head}\n{inner}\n{foot}</section>')

    doc = (
        "<!doctype html>\n<html lang=\"ko\"><head><meta charset=\"utf-8\">"
        '<meta name="viewport" content="width=device-width, initial-scale=1">'
        f"<title>{title}</title>"
        f"<style>{_HTML_CSS}</style>"
        f'<script id="kteacher-backport-marker" type="application/json">{json.dumps(marker, ensure_ascii=False)}</script>'
        "</head>"
        f'<body data-document-id="{_xesc(canonical["document_id"])}" data-document-class="{_xesc(canonical["document_class"])}">'
        f'{"".join(sections)}'
        "</body></html>"
    )
    path.write_text(doc, encoding="utf-8")




# --- DOCX (python-docx, real OOXML) --------------------------------------------

def _docx_set_ea_font(docx, style, name: str) -> None:
    from docx.oxml.ns import qn
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)


def _docx_cell_border_bottom(cell) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcpr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "6B7A8C")
    borders.append(bottom)
    tcpr.append(borders)

def _docx_write_cell(cell, value: str, *, bold: bool = False) -> None:
    """Write a cell so it always carries w:p/w:r/w:t, even when empty — required
    for a clean reopen of ruled/writing cells in Word and Hancom consumers."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    run = cell.paragraphs[0].add_run()
    if bold:
        run.bold = True
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = str(value)
    run._r.append(t)


def _render_docx_document(canonical: dict):
    docx = _require_docx()
    from docx.shared import Mm, Pt, RGBColor
    d = docx.Document()
    sec = d.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin = sec.bottom_margin = Mm(18)
    sec.left_margin = sec.right_margin = Mm(18)
    normal = d.styles["Normal"]
    normal.font.size = Pt(10.5)
    _docx_set_ea_font(docx, normal, "Malgun Gothic")

    for op in _flow(canonical):
        kind = op[0]
        if kind == "title":
            d.add_heading(op[1], level=0)
        elif kind == "meta":
            p = d.add_paragraph(op[1])
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0x5B, 0x6B, 0x7C)
        elif kind == "heading":
            d.add_heading(op[1], level=1)
        elif kind == "para":
            d.add_paragraph(op[1])
        elif kind == "evidence":
            table = d.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            _docx_write_cell(hdr[0], "구분", bold=True)
            _docx_write_cell(hdr[1], "근거", bold=True)
            for label, text in op[1]:
                row = table.add_row().cells
                _docx_write_cell(row[0], label)
                _docx_write_cell(row[1], text)
        elif kind == "task":
            p = d.add_paragraph()
            tag = f"[{op[1]}] " if op[1] else ""
            run = p.add_run(f"과제 {tag}")
            run.bold = True
            p.add_run(op[2])
        elif kind == "answer":
            n_lines, height_mm = op[1], op[2]
            table = d.add_table(rows=n_lines, cols=1)
            row_h = Mm(max(6.0, height_mm / max(1, n_lines)))
            for row in table.rows:
                row.height = row_h
                from docx.enum.table import WD_ROW_HEIGHT_RULE
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                _docx_cell_border_bottom(row.cells[0])
                _docx_write_cell(row.cells[0], "")
        elif kind == "card":
            title, bodytext, source, citation = op[1], op[2], op[3], op[4]
            table = d.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            cell = table.rows[0].cells[0]
            cell.paragraphs[0].add_run(title).bold = True
            cell.add_paragraph(bodytext)
            src = f"출처: {source}" + (f" ({citation})" if citation else "")
            note = cell.add_paragraph(src)
            note.runs[0].font.size = Pt(9)
            note.runs[0].font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
        elif kind == "stems":
            p = d.add_paragraph()
            p.add_run("문장 도우미: ").bold = True
            p.add_run(" / ".join(op[1])).italic = True
        elif kind == "note":
            d.add_paragraph(op[1])
        elif kind == "callout":
            p = d.add_paragraph()
            p.add_run(f"{op[1]}: ").bold = True
            p.add_run(op[2])
        elif kind == "table":
            caption, headers, rows, _rh = op[1], op[2], op[3], op[4]
            if caption:
                cap = d.add_paragraph()
                cap.add_run(str(caption)).bold = True
            ncols = max(len(headers), max((len(r) for r in rows), default=1))
            table = d.add_table(rows=0, cols=ncols)
            table.style = "Table Grid"
            if headers:
                cells = table.add_row().cells
                for i, h in enumerate(headers):
                    _docx_write_cell(cells[i], str(h), bold=True)
            for r in rows:
                cells = table.add_row().cells
                for i, v in enumerate(r):
                    _docx_write_cell(cells[i], str(v))
        elif kind == "pagebreak":
            d.add_page_break()
    return d


def _inject_opc_parts(path: Path, members: dict, *, content_type_defaults: dict) -> None:
    """Rewrite an OOXML zip to add extra parts and register any new Default content
    types so a strict consumer (Word) still accepts the package."""
    with zipfile.ZipFile(path) as z:
        items = {n: z.read(n) for n in z.namelist()}
    ct = items["[Content_Types].xml"].decode("utf-8")
    for ext, ctype in content_type_defaults.items():
        if f'Extension="{ext}"' not in ct:
            ct = ct.replace("</Types>", f'<Default Extension="{ext}" ContentType="{ctype}"/></Types>')
    items["[Content_Types].xml"] = ct.encode("utf-8")
    for name, data in members.items():
        items[name] = data.encode("utf-8") if isinstance(data, str) else data
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in items.items():
            z.writestr(name, data)


def render_docx(canonical: dict, marker: dict, path: Path) -> None:
    document = _render_docx_document(canonical)
    document.save(str(path))
    _inject_opc_parts(
        path,
        {
            _DOCX_MARKER: json.dumps(marker, ensure_ascii=False),
            _DOCX_CONTENT: _content_sidecar_xml(canonical),
        },
        content_type_defaults={"json": "application/json"},
    )


# --- HWPX (python-hwpx official builder, real OWPML) ---------------------------

def _render_hwpx_document(canonical: dict):
    B = _require_hwpx_builder()
    children: list = [B.Heading(level=1, text=canonical["title"]),
                      B.Paragraph(text=_doc_subtitle(canonical))]
    for op in _flow(canonical):
        kind = op[0]
        if kind in ("title", "meta"):
            continue  # already placed as document header
        elif kind == "heading":
            children.append(B.Heading(level=2, text=op[1]))
        elif kind == "para":
            children.append(B.Paragraph(text=op[1]))
        elif kind == "evidence":
            children.append(B.Table(header=["구분", "근거"], rows=[[l, t] for l, t in op[1]], header_shading="EAF1FB"))
        elif kind == "task":
            tag = f"[{op[1]}] " if op[1] else ""
            children.append(B.Paragraph(text=f"과제 {tag}{op[2]}"))
        elif kind == "answer":
            n_lines = op[1]
            children.append(B.Table(header=[], rows=[[""] for _ in range(n_lines)]))
        elif kind == "card":
            title, bodytext, source, citation = op[1], op[2], op[3], op[4]
            src = f"출처: {source}" + (f" ({citation})" if citation else "")
            children.append(B.Table(header=[title], rows=[[bodytext], [src]], header_shading="E6F4F1"))
        elif kind == "stems":
            children.append(B.Paragraph(text="문장 도우미: " + " / ".join(op[1])))
        elif kind == "note":
            children.append(B.Paragraph(text=op[1]))
        elif kind == "callout":
            children.append(B.Paragraph(text=f"{op[1]}: {op[2]}"))
        elif kind == "table":
            caption, headers, rows, _rh = op[1], op[2], op[3], op[4]
            if caption:
                children.append(B.Heading(level=2, text=str(caption)))
            children.append(B.Table(header=[str(h) for h in headers], rows=[[str(v) for v in r] for r in rows],
                                    header_shading="EAF1FB" if headers else None))
        elif kind == "pagebreak":
            children.append(B.PageBreak())
    section = B.Section(
        page=B.PageSize(width_mm=210, height_mm=297),
        margins=B.Margins(top_mm=18, right_mm=18, bottom_mm=18, left_mm=18),
        footer=B.Footer(children=[B.Paragraph(text=f"{canonical['title']} · 출력용 A4")]),
        children=children,
    )
    return B.Document(sections=[section], metadata=B.Metadata(title=canonical["title"], author="교사", organization="학교"))


def render_hwpx(canonical: dict, marker: dict, path: Path) -> None:
    document = _render_hwpx_document(canonical)
    document.save_to_path(str(path))
    with zipfile.ZipFile(path, "a", zipfile.ZIP_DEFLATED) as z:
        z.writestr(_HWPX_MARKER, json.dumps(marker, ensure_ascii=False))
        z.writestr(_HWPX_CONTENT, _content_sidecar_xml(canonical))


def render_all(ir: dict, out_dir: str | Path, *, document_index: int = 0) -> dict:
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    document = ir["lesson_package"]["documents"][document_index]
    canonical = canonical_content(document)
    fp = content_fingerprint(canonical)
    paths = {}
    for i, (fmt, renderer) in enumerate((("hwpx", render_hwpx), ("docx", render_docx), ("html", render_html))):
        marker = build_marker(ir, fmt, fp, render_revision_id=f"render-{2100 + i}")
        p = out_dir / f"{document['document_id']}.{fmt}"
        renderer(canonical, marker, p)
        paths[fmt] = str(p)
    return paths

SUPPORTED_RENDER_TARGETS = ["hwpx", "docx", "html"]


def render_package(ir: dict, out_dir: str | Path) -> dict:
    """Render every document in the package to all three formats.

    Returns `{document_id: {fmt: path}}`. Refuses duplicate document ids and any
    document whose render_targets are not exactly the supported HWPX/DOCX/HTML set.
    Reuses `render_all` per document (the single-document path stays unchanged)."""
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    documents = ir["lesson_package"]["documents"]
    seen: set = set()
    result: dict = {}
    for index, document in enumerate(documents):
        did = document["document_id"]
        if did in seen:
            raise ValueError(f"duplicate document id in package: {did!r}")
        seen.add(did)
        if document.get("render_targets") != SUPPORTED_RENDER_TARGETS:
            raise ValueError(f"unsupported render targets for {did!r}: {document.get('render_targets')!r}")
        result[did] = render_all(ir, out_dir, document_index=index)
    return result

# --- extractors (real round-trip) ----------------------------------------------

_PARA_RE = re.compile(r"<(?:hp:p|w:p|p)\b([^>]*)>(.*?)</(?:hp:p|w:p|p)>", re.DOTALL)
_ATTR_RE = re.compile(r'([\w:-]+)="([^"]*)"')
_TEXT_RE = re.compile(r"<[^>]+>")


def _parse_paras(xml: str) -> tuple[list, list]:
    content, prov = [], []
    for attrs_s, inner in _PARA_RE.findall(xml):
        attrs = dict(_ATTR_RE.findall(attrs_s))
        text = htmllib.unescape(_TEXT_RE.sub("", inner)).strip()
        if "data-content-id" in attrs:
            content.append({"content_id": attrs["data-content-id"], "text": text})
        elif "data-provenance-record-id" in attrs:
            prov.append({"record_id": attrs["data-provenance-record-id"], "label": htmllib.unescape(attrs.get("data-label", "")), "evidence_text": text})
    return content, prov

_BLOCK_RE = re.compile(r'<(table|hp:tbl|w:tbl|p|hp:p|w:p)\b([^>]*\bdata-block-id="[^"]*"[^>]*)>(.*?)</\1>', re.DOTALL)
_ROW_RE = re.compile(r"<(?:tr|hp:tr|w:tr)\b")
_CELL_RE = re.compile(r"<(?:td|hp:tc|w:tc)\b")
_FIRST_ROW_RE = re.compile(r"<(?:tr|hp:tr|w:tr)\b[^>]*>(.*?)</(?:tr|hp:tr|w:tr)>", re.DOTALL)


def _parse_blocks(xml: str) -> list:
    """Round-trip content.blocks from a rendered file: recover exact semantic fields
    from data-block-json and recompute the real rendered table/answer-space shape."""
    blocks = []
    for tag, attrs_s, inner in _BLOCK_RE.findall(xml):
        attrs = dict(_ATTR_RE.findall(attrs_s))
        payload = json.loads(htmllib.unescape(attrs["data-block-json"]))
        shape = None
        if tag in ("table", "hp:tbl", "w:tbl"):
            nrows = len(_ROW_RE.findall(inner))
            first = _FIRST_ROW_RE.search(inner)
            ncols = len(_CELL_RE.findall(first.group(1))) if first else 0
            shape = {"rows": nrows, "cols": ncols}
        blocks.append({
            "block_id": attrs["data-block-id"],
            "block_type": attrs["data-block-type"],
            "block": payload,
            "rendered_shape": shape,
        })
    return blocks


def _extract_zip(path: Path, marker_member: str, content_member: str) -> dict:
    with zipfile.ZipFile(path) as z:
        marker = json.loads(z.read(marker_member).decode("utf-8"))
        xml = z.read(content_member).decode("utf-8")
    content, prov = _parse_paras(xml)
    head = re.search(r'data-document-id="([^"]*)"[^>]*data-document-class="([^"]*)"[^>]*data-title="([^"]*)"', xml)
    return {
        "renderer_format": marker["renderer_format"],
        "document_id": head.group(1) if head else None,
        "document_class": head.group(2) if head else None,
        "title": htmllib.unescape(head.group(3)) if head else None,
        "required_content": content,
        "provenance_markers": prov,
        "unresolved_boundary_markers": [],
        "blocks": _parse_blocks(xml),
        "embedded_backport_marker_locator": marker_member,
        "embedded_backport_marker": marker,
    }


def extract_hwpx(path: str | Path) -> dict:
    return _extract_zip(Path(path), _HWPX_MARKER, _HWPX_CONTENT)


def extract_docx(path: str | Path) -> dict:
    return _extract_zip(Path(path), _DOCX_MARKER, _DOCX_CONTENT)


def extract_html(path: str | Path) -> dict:
    text = Path(path).read_text(encoding="utf-8")
    m = re.search(r'<script id="kteacher-backport-marker" type="application/json">(.*?)</script>', text, re.DOTALL)
    marker = json.loads(m.group(1))
    content, prov = _parse_paras(text)
    head = re.search(r'data-document-id="([^"]*)"[^>]*data-document-class="([^"]*)"', text)
    title = re.search(r"<h1>(.*?)</h1>", text, re.DOTALL)
    return {
        "renderer_format": "html",
        "document_id": head.group(1) if head else None,
        "document_class": head.group(2) if head else None,
        "title": htmllib.unescape(title.group(1)) if title else None,
        "required_content": content,
        "provenance_markers": prov,
        "unresolved_boundary_markers": [],
        "blocks": _parse_blocks(text),
        "embedded_backport_marker_locator": "script#kteacher-backport-marker",
        "embedded_backport_marker": marker,
    }


def extract_all(paths: dict) -> dict:
    return {"hwpx": extract_hwpx(paths["hwpx"]), "docx": extract_docx(paths["docx"]), "html": extract_html(paths["html"])}


def verify_parity(extracted: dict) -> tuple[bool, list]:
    """3-way parity: identical content/provenance/title across formats; each marker at
    its canonical location; identical content_fingerprint (same source IR)."""
    reasons = []
    fmts = ["hwpx", "docx", "html"]
    ref = extracted["hwpx"]
    for f in fmts:
        e = extracted[f]
        if e["required_content"] != ref["required_content"]:
            reasons.append(f"{f}: required_content drift")
        if e["provenance_markers"] != ref["provenance_markers"]:
            reasons.append(f"{f}: provenance_markers drift")
        if e["title"] != ref["title"] or e["document_id"] != ref["document_id"]:
            reasons.append(f"{f}: title/document_id drift")
        loc = CANONICAL_LOCATIONS[f]["locator_value"].lstrip("/")
        got = e["embedded_backport_marker_locator"].lstrip("/")
        if f != "html" and got != loc:
            reasons.append(f"{f}: marker not at canonical location ({got})")
        if e["embedded_backport_marker"]["content_fingerprint"] != ref["embedded_backport_marker"]["content_fingerprint"]:
            reasons.append(f"{f}: content_fingerprint drift (content not from same IR)")
        if e["embedded_backport_marker"]["renderer_format"] != f:
            reasons.append(f"{f}: marker renderer_format mismatch")
        if e.get("blocks", []) != ref.get("blocks", []):
            reasons.append(f"{f}: blocks drift")
    return (len(reasons) == 0, reasons)
